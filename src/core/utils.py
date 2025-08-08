
from __future__ import annotations
import re, io, time, os
from functools import wraps
from typing import Dict, List, Any, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
import pandas as pd
from markdown import markdown

from config_manager import API_RETRY_DELAYS, DEFAULT_STYLES

def retry_on_failure(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        attempts = len(API_RETRY_DELAYS) + 1
        for i in range(attempts):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                if i < len(API_RETRY_DELAYS):
                    time.sleep(API_RETRY_DELAYS[i])
                else:
                    raise
    return wrapper

@retry_on_failure
def call_openai(model_name: str, prompt: str, api_key: str,
                temperature: float = 0.7, top_p: float = 0.9, max_output_tokens: int = 1024) -> str:
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    # Utilise chat completions car responses.create n'existe pas dans l'API standard
    resp = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_output_tokens
    )
    if hasattr(resp, "choices") and resp.choices:
        return resp.choices[0].message.content or ""
    return str(resp)

@retry_on_failure
def call_anthropic(model_name: str, prompt: str, api_key: str,
                   temperature: float = 0.7, top_p: float = 0.9, max_output_tokens: int = 1024) -> str:
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=model_name,
        max_tokens=max_output_tokens,
        temperature=temperature,
        top_p=top_p,
        messages=[{"role":"user","content":prompt}]
    )
    parts = []
    for blk in getattr(msg, "content", []) or []:
        if isinstance(blk, dict) and blk.get("type") == "text":
            parts.append(blk.get("text",""))
        elif getattr(blk, "type", None) == "text":
            parts.append(getattr(blk, "text", ""))
    return "\n".join(parts).strip()

def parse_docx_plan(docx_path: str) -> List[Dict[str, Any]]:
    doc = Document(docx_path)
    heading_map = {"Heading 1":1,"Heading 2":2,"Heading 3":3,"Titre 1":1,"Titre 2":2,"Titre 3":3}
    counters = {1:0,2:0,3:0}
    items = []
    for p in doc.paragraphs:
        level = heading_map.get(getattr(getattr(p, "style", None), "name", ""))
        if level:
            for lv in [3,2,1]:
                if lv > level: counters[lv] = 0
            counters[level] += 1
            parts = [str(counters[1])] if counters[1] else []
            if counters[2]: parts.append(str(counters[2]))
            if counters[3]: parts.append(str(counters[3]))
            code = ".".join(parts) if parts else ""
            items.append({"code": code, "title": p.text.strip(), "level": level})
    return items

def _apply_doc_styles(doc, styles: dict):
    section = doc.sections[0]
    m = styles.get("margins_cm", DEFAULT_STYLES["margins_cm"])
    section.top_margin = Cm(m.get("top", 2.54))
    section.bottom_margin = Cm(m.get("bottom", 2.54))
    section.left_margin = Cm(m.get("left", 2.54))
    section.right_margin = Cm(m.get("right", 2.54))

def _add_paragraph(doc, text: str, size: int, font: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)

def generate_styled_docx(markdown_text: str, output_path: str, styles: Dict[str, Any]) -> None:
    font_family = styles.get("font_family", DEFAULT_STYLES["font_family"])
    body = int(styles.get("body_font_size", DEFAULT_STYLES["body_font_size"]))
    h1 = int(styles.get("h1_font_size", DEFAULT_STYLES["h1_font_size"]))
    h2 = int(styles.get("h2_font_size", DEFAULT_STYLES["h2_font_size"]))
    doc = Document()
    _apply_doc_styles(doc, styles)
    for line in markdown_text.splitlines():
        s = line.strip()
        if not s: 
            doc.add_paragraph("")
            continue
        if s.startswith("# "): _add_paragraph(doc, s[2:], h1, font_family); continue
        if s.startswith("## "): _add_paragraph(doc, s[3:], h2, font_family); continue
        if s.startswith(("- ", "* ")):
            p = doc.add_paragraph(s[2:]); p.style = "List Bullet"
            for r in p.runs: r.font.name = font_family; r.font.size = Pt(body)
            continue
        _add_paragraph(doc, s, body, font_family)
    doc.save(output_path)

def extract_used_references_apa(text_md: str) -> List[str]:
    pats = re.findall(r"\(([^,]+),\s*(\d{4})\)", text_md)
    return sorted({f"{a}, {y}" for a,y in pats})

def truncate_to_tokens(text: str, max_tokens: int, model: str = "gpt-4") -> str:
    """Tronque le texte selon le nombre max de tokens.
    Utilise tiktoken si disponible, sinon heuristique ~4 caractères = 1 token."""
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        tokens = enc.encode(text)
        if len(tokens) <= max_tokens:
            return text
        return enc.decode(tokens[:max_tokens])
    except Exception:
        # Fallback heuristique : ~4 caractères par token
        max_chars = max(256, int(max_tokens * 4))
        return text[:max_chars] if len(text) > max_chars else text

def _generate_filename(base_name: str, mode: str) -> str:
    """Génère un nom de fichier avec timestamp."""
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    # Nettoie le base_name pour éviter les caractères problématiques
    slug = "".join(c if c.isalnum() or c in "-_" else "-" for c in base_name)[:80]
    return f"{timestamp}_{mode}_{slug or 'sortie'}"

def export_markdown(text_md: str, base_name: str, mode: str, export_dir: str = "output") -> str:
    """Exporte le texte markdown dans un fichier .md."""
    outdir = Path(export_dir)
    outdir.mkdir(parents=True, exist_ok=True)
    filename = _generate_filename(base_name, mode)
    path = outdir / f"{filename}.md"
    path.write_text(text_md, encoding="utf-8")
    return str(path)

def export_docx(text_md: str, base_name: str, mode: str, export_dir: str = "output", styles: dict = None) -> str:
    """Exporte le texte markdown dans un fichier .docx stylé."""
    outdir = Path(export_dir)
    outdir.mkdir(parents=True, exist_ok=True)
    filename = _generate_filename(base_name, mode)
    path = outdir / f"{filename}.docx"
    generate_styled_docx(text_md, str(path), styles or {})
    return str(path)

def generate_bibliography(used: List[str], excel_path: str) -> str:
    import pandas as pd
    try:
        df = pd.read_excel(excel_path, sheet_name="Bibliographie")
    except Exception:
        df = pd.read_csv(excel_path)
    cols = {c.lower(): c for c in df.columns}
    col_full = cols.get("référence apa complète") or cols.get("reference apa complete") or cols.get("apa_full")
    col_short = cols.get("référence courte") or cols.get("reference courte") or cols.get("apa_short")
    if not col_full:
        return "\n".join(f"- {ref}" for ref in used)
    lines = []
    shorts = set(used)
    if col_short and col_short in df.columns:
        m = df[df[col_short].astype(str).str.strip().isin(shorts)]
        for _, r in m.iterrows(): lines.append(f"- {r[col_full]}")
    else:
        for ref in used:
            year = ref.split(",")[-1].strip()
            author = ref.rsplit(",",1)[0].strip()
            mask = df[col_full].astype(str).str.contains(author) & df[col_full].astype(str).str.contains(year)
            if mask.any(): lines.append(f"- {df[mask].iloc[0][col_full]}")
            else: lines.append(f"- {ref}")
    return "\n".join(lines)
